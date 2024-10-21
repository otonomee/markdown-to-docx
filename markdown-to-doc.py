import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys
import os

def create_word_styles(doc):
    # Name style
    name_style = doc.styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.size = Pt(18)
    name_style.font.bold = True
    name_style.paragraph_format.space_after = Pt(6)

    # Contact style
    contact_style = doc.styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.size = Pt(11)
    contact_style.paragraph_format.space_after = Pt(12)

    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(14)
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)

    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(12)
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.space_before = Pt(6)
    h2_style.paragraph_format.space_after = Pt(3)

    # Normal text style
    normal_style = doc.styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

def markdown_to_docx(md_file, docx_file):
    # Read Markdown content
    with open(md_file, 'r') as file:
        md_content = file.read()

    # Convert Markdown to HTML
    html = markdown.markdown(md_content)

    # Create a new Document
    doc = Document()
    create_word_styles(doc)

    # Parse HTML and add content to the document
    lines = html.split('\n')
    for line in lines:
        if line.startswith('<h1>'):
            p = doc.add_paragraph(re.sub('<[^<]+?>', '', line), style='Name')
        elif line.startswith('<h2>'):
            p = doc.add_paragraph(re.sub('<[^<]+?>', '', line), style='Heading 1')
        elif line.startswith('<h3>'):
            p = doc.add_paragraph(re.sub('<[^<]+?>', '', line), style='Heading 2')
        elif line.startswith('<p>'):
            text = re.sub('<[^<]+?>', '', line)
            if '|' in text:  # Assume this is contact information
                p = doc.add_paragraph(text, style='Contact')
            else:
                p = doc.add_paragraph(text, style='Normal')
        elif line.startswith('<ul>'):
            continue
        elif line.startswith('<li>'):
            p = doc.add_paragraph(re.sub('<[^<]+?>', '', line), style='Normal')
            p.style.paragraph_format.left_indent = Pt(18)
            p.style.paragraph_format.first_line_indent = Pt(-18)
            runs = p.runs
            runs[0].text = '• ' + runs[0].text

    # Save the document
    doc.save(docx_file)

def main():
    if len(sys.argv) != 2:
        print("Usage: python markdown_to_docx.py <input_markdown_file>")
        sys.exit(1)

    md_file = sys.argv[1]
    if not os.path.exists(md_file):
        print(f"Error: File '{md_file}' not found.")
        sys.exit(1)

    docx_file = os.path.splitext(md_file)[0] + '.docx'
    markdown_to_docx(md_file, docx_file)
    print(f"Converted '{md_file}' to '{docx_file}' successfully.")

if __name__ == "__main__":
    main()
